"""多格式文档解析路由：txt/md直读，docx用python-docx，pdf用pymupdf，扫描件路由OCR分支。"""
from pathlib import Path


class UnsupportedFormat(Exception):
    pass


def parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise UnsupportedFormat("无法识别文本编码")


def parse_docx(data: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n\n".join(parts)


def parse_pdf(data: bytes) -> tuple[str, str]:
    """返回 (文本, 通道)。无文本层的扫描件路由至OCR通道。"""
    import fitz  # pymupdf
    doc = fitz.open(stream=data, filetype="pdf")
    text = "\n\n".join(page.get_text() for page in doc)
    if len(text.strip()) < 20:
        # 扫描件：生产环境调用PaddleOCR服务，demo环境降级提示
        return ocr_fallback(data), "ocr"
    return text, "text-layer"


def ocr_fallback(data: bytes) -> str:
    """OCR通道占位实现。生产环境对接 PaddleOCR / 腾讯云OCR API。"""
    raise UnsupportedFormat("已路由至OCR通道：demo环境未加载OCR模型，生产环境使用PaddleOCR处理扫描件")


def parse_any(filename: str, data: bytes) -> tuple[str, str]:
    """格式嗅探并解析，返回 (文本, 解析通道)。"""
    suffix = Path(filename).suffix.lower()
    if suffix in (".txt", ".md"):
        return parse_txt(data), "text"
    if suffix == ".docx":
        return parse_docx(data), "docx"
    if suffix == ".pdf" or data[:4] == b"%PDF":
        return parse_pdf(data)
    raise UnsupportedFormat(f"暂不支持的格式: {suffix or '未知'}（支持 txt/md/docx/pdf）")
